#!/usr/bin/env python3
"""Экспорт JSON-отчёта парсера в один DOCX."""

from __future__ import annotations

import argparse
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

PROJECT_CHAT_RE = re.compile(
    r"(actinsta|instachat|instab|instacyber|instal|likebunny|likecosmo|squanch|actdino|mpickles)",
    re.I,
)


def is_project_entity(entity: dict) -> bool:
    username = (entity.get("username") or "").lower()
    title = (entity.get("title") or "").lower()
    if PROJECT_CHAT_RE.search(username) or PROJECT_CHAT_RE.search(title):
        return True
    if entity.get("entity_type") in {"bot"} and username in {
        "actdino_bot",
        "mpickles_bot",
    }:
        return True
    return entity.get("depth", 99) <= 1 and entity.get("entity_type") in {
        "supergroup",
        "channel",
        "bot",
    }


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_kv(doc: Document, key: str, value: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{key}: ")
    run.bold = True
    p.add_run(value or "—")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val[:5000] if val else ""
    doc.add_paragraph()


def truncate(text: str, limit: int = 1500) -> str:
    text = (text or "").strip()
    if len(text) <= limit:
        return text
    return text[:limit] + "…"


def build_doc(data: dict) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(data.get("project_name", "Telegram project"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_kv(doc, "Seed URL", data.get("seed_url", ""))
    add_kv(doc, "Дата сбора", data.get("crawled_at", ""))
    add_kv(doc, "Всего сущностей", str(data.get("total_entities", 0)))
    add_kv(doc, "Чатов/каналов", str(len(data.get("chats_and_channels", []))))
    add_kv(doc, "Ботов", str(len(data.get("bots", []))))
    add_kv(doc, "Уникальных ссылок", str(len(data.get("all_links", []))))

    entities = data.get("entities", [])
    project_entities = [e for e in entities if is_project_entity(e)]
    other_entities = [e for e in entities if not is_project_entity(e)]

    add_heading(doc, "1. Сводка по проекту конкурента", 1)
    doc.add_paragraph(
        "Ниже — чаты, каналы и боты сети Instagram-активности (actinsta / instachat и связанные). "
        "Остальные сущности (упоминания из сообщений участников) вынесены в конец."
    )

    add_heading(doc, "1.1 Чаты и каналы сети", 2)
    rows = []
    for e in sorted(
        project_entities,
        key=lambda x: (-(x.get("members_count") or 0), x.get("title", "")),
    ):
        if e.get("entity_type") not in {"supergroup", "channel", "invite_preview"}:
            continue
        rows.append(
            [
                e.get("title", ""),
                e.get("username", ""),
                e.get("url", ""),
                str(e.get("members_count") or ""),
                e.get("entity_type", ""),
            ]
        )
    add_table(doc, ["Название", "Username", "Ссылка", "Участники", "Тип"], rows)

    add_heading(doc, "1.2 Боты", 2)
    bot_rows = []
    for e in sorted(project_entities, key=lambda x: x.get("username", "")):
        if not e.get("is_bot"):
            continue
        bot_rows.append([e.get("title", ""), e.get("username", ""), e.get("url", "")])
    add_table(doc, ["Название", "Username", "Ссылка"], bot_rows)

    add_heading(doc, "1.3 Ключевые ссылки из сети", 2)
    key_links = sorted(
        {
            link
            for e in project_entities
            for link in e.get("links_found", [])
            if "t.me" in link and PROJECT_CHAT_RE.search(link)
        }
    )
    for link in key_links:
        doc.add_paragraph(link, style="List Bullet")

    add_heading(doc, "2. Детали по сущностям проекта", 1)
    for e in sorted(project_entities, key=lambda x: (x.get("depth", 0), x.get("title", ""))):
        add_heading(doc, f"{e.get('title') or e.get('username')} (@{e.get('username', '—')})", 2)
        add_kv(doc, "Ссылка", e.get("url", ""))
        add_kv(doc, "Тип", e.get("entity_type", ""))
        add_kv(doc, "Участники", str(e.get("members_count") or "—"))
        add_kv(doc, "Описание", truncate(e.get("description", ""), 3000))
        add_kv(doc, "Глубина обхода", str(e.get("depth", "")))
        add_kv(doc, "Сообщений просмотрено", str(e.get("messages_scanned", 0)))

        if e.get("links_found"):
            add_heading(doc, "Ссылки", 3)
            for link in e["links_found"][:50]:
                doc.add_paragraph(link, style="List Bullet")
            if len(e["links_found"]) > 50:
                doc.add_paragraph(f"… и ещё {len(e['links_found']) - 50} ссылок")

        if e.get("inline_buttons"):
            add_heading(doc, "Кнопки в сообщениях", 3)
            for btn in e["inline_buttons"][:30]:
                text = btn.get("text", "")
                url = btn.get("url", "")
                doc.add_paragraph(f"{text} → {url}" if url else text, style="List Bullet")

        if e.get("bot_replies"):
            add_heading(doc, "Ответы бота", 3)
            for reply in e["bot_replies"]:
                doc.add_paragraph(f"Команда: {reply.get('trigger', '')}")
                doc.add_paragraph(truncate(reply.get("text", ""), 2000))

        if e.get("sample_messages"):
            add_heading(doc, "Примеры сообщений", 3)
            for msg in e["sample_messages"][:8]:
                if isinstance(msg, dict):
                    text = msg.get("text", "")
                else:
                    text = str(msg)
                if text:
                    doc.add_paragraph(truncate(text, 1200))

        if e.get("errors"):
            add_kv(doc, "Ошибки", "; ".join(e["errors"]))

    add_heading(doc, "3. Все найденные ссылки", 1)
    for link in sorted(data.get("all_links", [])):
        if link.startswith("http"):
            doc.add_paragraph(link, style="List Bullet")

    add_heading(doc, "4. Прочие сущности (из сообщений участников)", 1)
    doc.add_paragraph(
        f"Всего: {len(other_entities)}. Чаще это @username из постов участников, "
        "которые не являются публичными каналами/чатами проекта."
    )
    other_rows = []
    for e in sorted(other_entities, key=lambda x: x.get("username", ""))[:200]:
        other_rows.append(
            [
                e.get("username", ""),
                e.get("title", ""),
                e.get("entity_type", ""),
                str(e.get("members_count") or ""),
                truncate("; ".join(e.get("errors", [])), 200),
            ]
        )
    add_table(
        doc,
        ["Username", "Название", "Тип", "Участники", "Ошибка/примечание"],
        other_rows,
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"Документ сформирован: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return doc


def main() -> None:
    p = argparse.ArgumentParser(description="Экспорт JSON в DOCX")
    p.add_argument("--input", default="output/instachat6_20260623_202001.json")
    p.add_argument("--output", default="")
    args = p.parse_args()

    input_path = Path(args.input)
    data = json.loads(input_path.read_text(encoding="utf-8"))

    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.with_name(input_path.stem + "_report.docx")

    doc = build_doc(data)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Сохранено: {output_path.resolve()}")


if __name__ == "__main__":
    main()
