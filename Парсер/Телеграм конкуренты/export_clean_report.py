#!/usr/bin/env python3
"""Чистый отчёт: только инфраструктура конкурента (чаты, каналы, боты)."""

from __future__ import annotations

import argparse
import csv
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Только сущности сети actinsta / instachat
CLEAN_USERNAME_RE = re.compile(
    r"^("
    r"actinsta\d*|"
    r"actinsta|"
    r"instachat\d*|"
    r"instabochat|"
    r"actinstchat|"
    r"instacyberchat|"
    r"squanch_chat|"
    r"likebunny|"
    r"likecosmo|"
    r"instalps\d+|"
    r"instals\d+|"
    r"actdino_bot|"
    r"mpickles_bot"
    r")$",
    re.I,
)

PROJECT_LINK_RE = re.compile(
    r"t\.me/(actinsta|instachat|instab|actinst|instacyber|squanch|likebunny|likecosmo|instal|instals|actdino|mpickles)",
    re.I,
)


def is_clean_entity(entity: dict) -> bool:
    username = (entity.get("username") or "").strip()
    if not username:
        return False
    if not CLEAN_USERNAME_RE.match(username):
        return False
    return entity.get("entity_type") in {"supergroup", "channel", "bot"}


def collect_unique_menu_buttons(entities: list[dict]) -> list[dict[str, str]]:
    """Один набор кнопок меню (без дублей по тексту)."""
    seen: set[str] = set()
    buttons: list[dict[str, str]] = []
    for entity in entities:
        for btn in entity.get("inline_buttons", []):
            text = (btn.get("text") or "").strip()
            url = (btn.get("url") or "").strip()
            if not text or not url:
                continue
            if not PROJECT_LINK_RE.search(url) and "mpickles" not in url and "actdino" not in url:
                continue
            if text in seen:
                continue
            seen.add(text)
            buttons.append({"text": text, "url": url})
    order = ["чаты", "vip", "правила", "список", "обратная", "связь"]
    def sort_key(btn: dict[str, str]) -> tuple[int, str]:
        low = btn["text"].lower()
        for i, word in enumerate(order):
            if word in low:
                return (i, btn["text"])
        return (99, btn["text"])
    return sorted(buttons, key=sort_key)


def filter_project_links(links: list[str]) -> list[str]:
    result = []
    for link in links or []:
        if PROJECT_LINK_RE.search(link):
            result.append(link)
    return sorted(set(result))


def load_pricing(pricing_path: Path | None) -> dict | None:
    if not pricing_path or not pricing_path.exists():
        return None
    try:
        return json.loads(pricing_path.read_text(encoding="utf-8"))
    except Exception:
        return None


def add_pricing_section(doc: Document, pricing: dict | None, ads: dict | None) -> None:
    doc.add_heading("4. Цены и тарифы", 1)
    doc.add_paragraph(
        "Источник: бот @mpickles_bot → «Условия и цены 💎 VIP 💎». "
        "Оплата VIP и рекламы — через администратора (6:00–20:00 МСК, @actinsta_admin)."
    )

    doc.add_heading("Бесплатный уровень", 2)
    doc.add_paragraph(
        "Взаимная активность в чатах сети: пользователь выполняет задания бота @actdino_bot "
        "и публикует свою ссылку в чат. Канал @actinsta позиционирует сеть как бесплатную."
    )

    doc.add_heading("VIP — стандартные тарифы (30 постов)", 2)
    vip_table = doc.add_table(rows=1, cols=4)
    vip_table.style = "Table Grid"
    for i, h in enumerate(["Тариф", "Состав", "Цена", "Ожидаемая активность"]):
        vip_table.rows[0].cells[i].text = h
    for row_data in [
        ("ЛС", "Лайк + сохранение", "1 200 ₽", "60–70 лайков/пост"),
        ("ЛПС", "Лайк + подписка + сохранение", "2 000 ₽", "60–70 лайков/пост, ~200–250 подписок суммарно"),
        ("ЛКС", "Лайк + комментарий + сохранение", "1 800 ₽", "20–30 лайков и комментариев/пост"),
        ("ЛКСП", "Лайк + коммент + сохран + подписка", "2 300 ₽", "20–30 лайков и комментов/пост, ~200–250 подписок"),
        ("Индивидуальный", "Свой объём активности", "по договорённости", "—"),
    ]:
        row = vip_table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_paragraph()
    doc.add_heading("Разовый заказ", 2)
    doc.add_paragraph("Комментарии (5 ₽ за комментарий):")
    comments_table = doc.add_table(rows=1, cols=2)
    comments_table.style = "Table Grid"
    comments_table.rows[0].cells[0].text = "Количество"
    comments_table.rows[0].cells[1].text = "Цена"
    for qty, price in [("10", "50 ₽"), ("20", "100 ₽"), ("30", "150 ₽"), ("50", "250 ₽"), ("100", "500 ₽")]:
        row = comments_table.add_row().cells
        row[0].text = qty
        row[1].text = price

    doc.add_paragraph()
    doc.add_paragraph("Лайки (1 ₽ за лайк):")
    likes_table = doc.add_table(rows=1, cols=2)
    likes_table.style = "Table Grid"
    likes_table.rows[0].cells[0].text = "Количество"
    likes_table.rows[0].cells[1].text = "Цена"
    for qty, price in [
        ("50", "50 ₽"),
        ("100", "100 ₽"),
        ("150", "150 ₽"),
        ("200", "200 ₽"),
        ("300", "300 ₽"),
    ]:
        row = likes_table.add_row().cells
        row[0].text = qty
        row[1].text = price

    doc.add_paragraph()
    doc.add_heading("Условия VIP", 2)
    for line in [
        "Не нужно выполнять задания других участников.",
        "Не нужно самостоятельно отправлять ссылку в чат — бот ставит её в очередь с другими VIP.",
        "Ссылку можно менять или останавливать с 6:00 до 20:00 МСК.",
        "Активность нарастает постепенно; в выходные может быть ниже.",
        "Участники — живые люди, правила проверяются ежедневно.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_paragraph()
    doc.add_heading("Реклама", 2)
    ads_text = "Публичного прайса нет. Оформление через бота → «Реклама 📣» → связь с администратором."
    if ads:
        for section in ads.get("sections", []):
            if section.get("button") == "Реклама 📣":
                for sub in section.get("subclicks", []):
                    if sub.get("button") == "Купить рекламу 🤝" and sub.get("text"):
                        ads_text = sub["text"].replace("**", "").replace("__", "")
                        break
    doc.add_paragraph(ads_text)

    doc.add_paragraph()
    doc.add_heading("Соответствие чатов типам активности", 2)
    chat_table = doc.add_table(rows=1, cols=2)
    chat_table.style = "Table Grid"
    chat_table.rows[0].cells[0].text = "Чат"
    chat_table.rows[0].cells[1].text = "Тип (бесплатно)"
    for chat, kind in [
        ("instachat6, actinstchat, squanch_chat, instabochat", "ЛКС 6"),
        ("actinsta1", "ЛКС 5"),
        ("instacyberchat", "ЛКСП 6"),
        ("instalps10", "ЛПС 10"),
        ("likebunny", "ЛС 10"),
        ("actinsta3", "ЛС 14"),
        ("likecosmo", "ЛС 15"),
        ("instals20", "ЛС 20"),
        ("actinsta2", "Слово/смайл ЛКС 6"),
    ]:
        row = chat_table.add_row().cells
        row[0].text = chat
        row[1].text = kind

    if pricing and pricing.get("vip_overview"):
        doc.add_paragraph()
        doc.add_heading("Текст из бота (оригинал)", 2)
        raw = pricing["vip_overview"]
        for marker in ("🔝 Главное Меню", "Привет! Чем я могу помочь"):
            if marker in raw:
                raw = raw.split(marker)[0]
        clean = raw.replace("**", "").replace("__", "").strip()
        doc.add_paragraph(clean)


def build_clean_doc(
    data: dict,
    entities: list[dict],
    *,
    pricing: dict | None = None,
    ads: dict | None = None,
) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Сеть конкурента: actinsta / instachat", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Seed: ").bold = True
    p.add_run(data.get("seed_url", ""))
    p = doc.add_paragraph()
    p.add_run("Дата сбора: ").bold = True
    p.add_run(data.get("crawled_at", ""))
    p = doc.add_paragraph()
    p.add_run("Чатов/каналов в сети: ").bold = True
    p.add_run(str(sum(1 for e in entities if e.get("entity_type") in {"supergroup", "channel"})))
    p = doc.add_paragraph()
    p.add_run("Ботов: ").bold = True
    p.add_run(str(sum(1 for e in entities if e.get("is_bot"))))

    doc.add_heading("1. Чаты и каналы", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Название", "Username", "Ссылка", "Участники", "Тип"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    for e in sorted(
        entities,
        key=lambda x: (-(x.get("members_count") or 0), x.get("title", "")),
    ):
        if e.get("entity_type") not in {"supergroup", "channel"}:
            continue
        row = table.add_row().cells
        row[0].text = e.get("title", "")
        row[1].text = f"@{e.get('username', '')}"
        row[2].text = e.get("url", "")
        row[3].text = str(e.get("members_count") or "—")
        row[4].text = e.get("entity_type", "")

    doc.add_paragraph()
    doc.add_heading("2. Боты", 1)
    bot_table = doc.add_table(rows=1, cols=3)
    bot_table.style = "Table Grid"
    for i, h in enumerate(["Название", "Username", "Ссылка"]):
        bot_table.rows[0].cells[i].text = h
    for e in sorted(entities, key=lambda x: x.get("username", "")):
        if not e.get("is_bot"):
            continue
        row = bot_table.add_row().cells
        row[0].text = e.get("title", "")
        row[1].text = f"@{e.get('username', '')}"
        row[2].text = e.get("url", "")

    menu_buttons = collect_unique_menu_buttons(entities)
    if menu_buttons:
        doc.add_paragraph()
        doc.add_heading("3. Типовое меню чата (один набор)", 1)
        doc.add_paragraph(
            "Такие кнопки закреплены в чатах сети. Показан один экземпляр без повторов."
        )
        for btn in menu_buttons:
            doc.add_paragraph(f"{btn['text']} → {btn['url']}")

    add_pricing_section(doc, pricing, ads)

    doc.add_paragraph()
    doc.add_heading("5. Карточки сущностей", 1)

    for e in sorted(entities, key=lambda x: (x.get("depth", 0), -(x.get("members_count") or 0))):
        doc.add_heading(f"{e.get('title', '')} (@{e.get('username', '')})", 2)
        for key, label in [
            ("url", "Ссылка"),
            ("entity_type", "Тип"),
            ("members_count", "Участники"),
            ("description", "Описание"),
            ("depth", "Глубина"),
        ]:
            val = e.get(key, "")
            if key == "members_count" and not val:
                val = "—"
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(str(val) if val else "—")

        links = filter_project_links(e.get("links_found", []))
        if links:
            doc.add_heading("Ссылки сети", 3)
            for link in links:
                doc.add_paragraph(link, style="List Bullet")

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.add_run(f"Сформировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return doc


def save_clean_csv(entities: list[dict], path: Path) -> None:
    fields = [
        "url",
        "username",
        "entity_type",
        "title",
        "description",
        "members_count",
        "is_bot",
        "depth",
        "links_found",
    ]
    with path.open("w", encoding="utf-8-sig", newline="") as fh:
        writer = csv.DictWriter(fh, fieldnames=fields, delimiter=";")
        writer.writeheader()
        for e in sorted(entities, key=lambda x: (-(x.get("members_count") or 0), x.get("title", ""))):
            writer.writerow(
                {
                    "url": e.get("url", ""),
                    "username": e.get("username", ""),
                    "entity_type": e.get("entity_type", ""),
                    "title": e.get("title", ""),
                    "description": (e.get("description") or "").replace("\n", " ")[:500],
                    "members_count": e.get("members_count") or "",
                    "is_bot": "yes" if e.get("is_bot") else "no",
                    "depth": e.get("depth", ""),
                    "links_found": " | ".join(filter_project_links(e.get("links_found", []))),
                }
            )


def main() -> None:
    p = argparse.ArgumentParser(description="Чистый отчёт без участников и сообщений")
    p.add_argument("--input", default="output/instachat6_20260623_202001.json")
    p.add_argument("--output", default="")
    p.add_argument("--pricing", default="output/pricing_complete.json")
    p.add_argument("--ads", default="output/pricing_full.json")
    args = p.parse_args()

    input_path = Path(args.input)
    data = json.loads(input_path.read_text(encoding="utf-8"))
    entities = [e for e in data.get("entities", []) if is_clean_entity(e)]

    stem = input_path.stem + "_clean"
    out_dir = input_path.parent
    docx_path = Path(args.output) if args.output else out_dir / f"{stem}_report_v3.docx"
    csv_path = out_dir / f"{stem}.csv"

    pricing = load_pricing(Path(args.pricing) if args.pricing else None)
    ads = load_pricing(Path(args.ads) if args.ads else None)
    doc = build_clean_doc(data, entities, pricing=pricing, ads=ads)
    doc.save(docx_path)
    save_clean_csv(entities, csv_path)

    print(f"Сущностей в чистом отчёте: {len(entities)}")
    print(f"DOCX: {docx_path.resolve()}")
    print(f"CSV:  {csv_path.resolve()}")


if __name__ == "__main__":
    main()
