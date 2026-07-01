"""Формирование подробного DOCX-отчёта."""

from __future__ import annotations

from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

SOURCE_LABELS = {
    "google": "Google",
    "yandex": "Яндекс",
    "vk": "VK",
    "telegram": "Telegram",
}


def _truncate(text: str | None, limit: int = 4000) -> str:
    text = (text or "").strip()
    if len(text) <= limit:
        return text
    return text[:limit] + "…"


def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _kv(doc: Document, key: str, value: str | int | None) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{key}: ")
    r.bold = True
    p.add_run(str(value) if value not in (None, "") else "—")


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(_truncate(text, 2000), style="List Bullet")


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not rows:
        doc.add_paragraph("Нет данных.")
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = _truncate(str(val) if val is not None else "", 3000)
    doc.add_paragraph()


def _counter_section(doc: Document, title: str, counter: dict[str, int]) -> None:
    _heading(doc, title, 2)
    if not counter:
        doc.add_paragraph("Нет данных.")
        return
    rows = [[k, str(v)] for k, v in sorted(counter.items(), key=lambda x: -x[1])]
    _table(doc, ["Показатель", "Количество"], rows)


def build_report_docx(data: dict[str, Any]) -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ── Титул ──
    title = doc.add_heading("Отчёт конкурентной разведки", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("AI-поисковик конкурентов · Verdi Monitor")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()
    _kv(doc, "Дата формирования", data.get("generated_at"))
    stats = data.get("stats", {})
    _kv(doc, "Всего находок", stats.get("total_findings", 0))
    _kv(doc, "С AI-анализом", stats.get("analyzed_findings", 0))
    _kv(doc, "Конкурентов в мониторинге", stats.get("competitors_count", 0))
    doc.add_page_break()

    # ── 1. Исполнительное резюме ──
    _heading(doc, "1. Исполнительное резюме", 1)
    if data.get("ai_executive_summary"):
        for para in data["ai_executive_summary"].split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    else:
        doc.add_paragraph(
            "AI-сводка недоступна (не настроен API или анализ не выполнен). "
            "Ниже представлена полная статистика и детализация по всем находкам."
        )
        doc.add_paragraph()
        doc.add_paragraph(
            f"Система зафиксировала {stats.get('total_findings', 0)} результатов "
            f"по {stats.get('competitors_count', 0)} конкурентам. "
            f"Основные источники: {', '.join(f'{k} ({v})' for k, v in stats.get('by_source', {}).items())}."
        )

    # ── 2. Полная картина рынка ──
    _heading(doc, "2. Полная картина рынка и конкурентной среды", 1)
    if data.get("ai_market_picture"):
        for para in data["ai_market_picture"].split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    else:
        doc.add_paragraph(
            "Раздел сформирован на основе агрегированных данных мониторинга."
        )

    _heading(doc, "2.1 Распределение по источникам", 2)
    rows = [[k, str(v)] for k, v in sorted(stats.get("by_source", {}).items(), key=lambda x: -x[1])]
    _table(doc, ["Источник", "Количество"], rows)

    _heading(doc, "2.2 Распределение по конкурентам", 2)
    rows = [[k, str(v)] for k, v in sorted(stats.get("by_competitor", {}).items(), key=lambda x: -x[1])]
    _table(doc, ["Конкурент", "Количество"], rows)

    _heading(doc, "2.3 Матрица: конкурент × источник", 2)
    matrix = data.get("by_competitor_source", {})
    if matrix:
        sources = sorted({s for cols in matrix.values() for s in cols})
        headers = ["Конкурент"] + [SOURCE_LABELS.get(s, s) for s in sources]
        rows = []
        for comp, cols in sorted(matrix.items()):
            rows.append([comp] + [str(cols.get(s, 0)) for s in sources])
        _table(doc, headers, rows)

    doc.add_page_break()

    # ── 3. Профили конкурентов ──
    _heading(doc, "3. Профили отслеживаемых конкурентов", 1)
    for c in data.get("competitors", []):
        _heading(doc, f"3.{c['id']}. {c['name']}", 2)
        _kv(doc, "Статус", "Активен" if c.get("is_active") else "Отключён")
        _kv(doc, "Регион", c.get("region"))
        _kv(doc, "Добавлен", c.get("created_at"))
        _kv(doc, "Находок", c.get("findings_count", 0))
        _kv(doc, "Брендовые ключи", _join_list(c.get("brand_keywords")))
        _kv(doc, "Коммерческие ключи", _join_list(c.get("money_keywords")))
        _kv(doc, "Google-запросы", _join_list(c.get("google_queries")))
        _kv(doc, "Яндекс-запросы", _join_list(c.get("yandex_queries")))
        _kv(doc, "VK-домены", _join_list(c.get("vk_domains")))
        _kv(doc, "VK owner_id", _join_list(c.get("vk_owner_ids")))
        _kv(doc, "Telegram-каналы", _join_list(c.get("telegram_channels")))
        doc.add_paragraph()

    doc.add_page_break()

    # ── 4. Аналитика AI ──
    _heading(doc, "4. Сводная AI-аналитика", 1)
    _counter_section(doc, "4.1 Тональность", stats.get("by_tone", {}))
    _counter_section(doc, "4.2 Сентимент", stats.get("by_sentiment", {}))
    _counter_section(doc, "4.3 Намерение (intent)", stats.get("by_intent", {}))
    _counter_section(doc, "4.4 Типы результатов", stats.get("by_type", {}))

    agg = data.get("aggregates", {})
    _heading(doc, "4.5 Сводка коммерческих офферов", 2)
    for item in agg.get("offers", []) or ["—"]:
        _bullet(doc, item)

    _heading(doc, "4.6 Призывы к действию (CTA)", 2)
    for item in agg.get("ctas", []) or ["—"]:
        _bullet(doc, item)

    _heading(doc, "4.7 Боли аудитории", 2)
    for item in agg.get("pain_points", []) or ["—"]:
        _bullet(doc, item)

    _heading(doc, "4.8 Ключевые триггеры и хуки", 2)
    for item in agg.get("hooks", []) or ["—"]:
        _bullet(doc, item)

    doc.add_page_break()

    # ── 5. Детальный разбор по источникам ──
    _heading(doc, "5. Детальный разбор по источникам", 1)
    findings = data.get("findings", [])
    for source_key, source_label in SOURCE_LABELS.items():
        source_findings = [f for f in findings if f["source"] == source_key]
        if not source_findings:
            continue
        _heading(doc, f"5.{source_key}. {source_label} ({len(source_findings)} находок)", 2)
        for f in source_findings:
            _add_finding_block(doc, f, level=3)

    doc.add_page_break()

    # ── 6. Детальный разбор по конкурентам ──
    _heading(doc, "6. Детальный разбор по конкурентам", 1)
    by_comp: dict[str, list] = {}
    for f in findings:
        by_comp.setdefault(f["competitor"], []).append(f)
    for comp_name, comp_findings in sorted(by_comp.items()):
        _heading(doc, f"{comp_name} ({len(comp_findings)} находок)", 2)
        for f in comp_findings:
            _add_finding_block(doc, f, level=3)

    doc.add_page_break()

    # ── 7. Полный каталог находок ──
    _heading(doc, "7. Полный каталог находок", 1)
    doc.add_paragraph(
        f"Исчерпывающий перечень всех {len(findings)} записей с полными текстами и метаданными."
    )
    rows = []
    for f in findings:
        a = f.get("analysis") or {}
        rows.append(
            [
                str(f["id"]),
                f["competitor"],
                SOURCE_LABELS.get(f["source"], f["source"]),
                f["result_type"],
                _truncate(f.get("title") or f.get("raw_text"), 120),
                a.get("tone") or "—",
                a.get("sentiment") or "—",
                f.get("url") or "—",
                f.get("collected_at", "—"),
            ]
        )
    _table(
        doc,
        ["ID", "Конкурент", "Источник", "Тип", "Заголовок/текст", "Tone", "Sentiment", "URL", "Собрано"],
        rows,
    )

    for f in findings:
        _heading(doc, f"Находка #{f['id']}", 2)
        _add_finding_block(doc, f, level=3, full=True)

    doc.add_page_break()

    # ── 8. Рекомендации ──
    _heading(doc, "8. Рекомендации и следующие шаги", 1)
    if data.get("ai_recommendations"):
        for para in data["ai_recommendations"].split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    else:
        doc.add_paragraph("Рекомендации будут сформированы после настройки AI API и повторного экспорта.")
        _bullet(doc, "Запустите AI-анализ находок (ЗАПУСТИТЬ_АНАЛИЗ.bat)")
        _bullet(doc, "Добавьте ключи Google / Яндекс / VK для расширения охвата")
        _bullet(doc, "Настройте расписание фонового мониторинга в Settings")

    # ── 9. Методология ──
    _heading(doc, "9. Методология и параметры мониторинга", 1)
    settings = data.get("settings", {})
    _kv(doc, "AI-модель", settings.get("ai_model"))
    _kv(doc, "Интервал мониторинга (часы)", settings.get("monitor_interval_hours"))
    _kv(doc, "Google", "вкл." if settings.get("google_enabled") else "выкл.")
    _kv(doc, "Яндекс", "вкл." if settings.get("yandex_enabled") else "выкл.")
    _kv(doc, "VK", "вкл." if settings.get("vk_enabled") else "выкл.")
    _kv(doc, "Telegram", "вкл." if settings.get("telegram_enabled") else "выкл.")
    doc.add_paragraph()
    doc.add_paragraph(
        "Источники данных: Google Custom Search, Yandex Search API, VK API (wall.get, search), "
        "Telegram (Telethon). AI-анализ выполняется через OpenAI-совместимый API. "
        "Дедупликация по source+external_id или hash текста/URL."
    )

    footer = doc.add_paragraph()
    footer.add_run(f"Сформировано Verdi Monitor · {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    return doc


def _join_list(items: list | None) -> str:
    if not items:
        return "—"
    return ", ".join(str(x) for x in items)


def _add_finding_block(doc: Document, f: dict, level: int = 3, full: bool = False) -> None:
    title = f.get("title") or _truncate(f.get("raw_text"), 80) or f"Находка #{f['id']}"
    _heading(doc, title, level)
    _kv(doc, "ID", f["id"])
    _kv(doc, "Конкурент", f.get("competitor"))
    _kv(doc, "Источник", SOURCE_LABELS.get(f["source"], f["source"]))
    _kv(doc, "Тип", f.get("result_type"))
    _kv(doc, "Канал / автор", f.get("channel_name") or f.get("author_name"))
    _kv(doc, "URL", f.get("url"))
    _kv(doc, "Позиция в выдаче", f.get("position"))
    _kv(doc, "Просмотры", f.get("views"))
    _kv(doc, "Лайки", f.get("likes"))
    _kv(doc, "Репосты", f.get("reposts"))
    _kv(doc, "Комментарии", f.get("comments"))
    _kv(doc, "Дата публикации", f.get("published_at"))
    _kv(doc, "Дата сбора", f.get("collected_at"))

    if f.get("snippet"):
        _kv(doc, "Сниппет", _truncate(f["snippet"], 2000 if full else 500))

    if f.get("raw_text"):
        _heading(doc, "Полный текст", level + 1)
        doc.add_paragraph(_truncate(f["raw_text"], 8000 if full else 2000))

    a = f.get("analysis")
    if a:
        _heading(doc, "AI-анализ", level + 1)
        _kv(doc, "Тип сущности", a.get("entity_type"))
        _kv(doc, "Резюме", a.get("summary"))
        _kv(doc, "Оффер", a.get("offer"))
        _kv(doc, "CTA", a.get("cta"))
        _kv(doc, "Тональность", a.get("tone"))
        _kv(doc, "Сентимент", a.get("sentiment"))
        _kv(doc, "Намерение", a.get("intent"))
        _kv(doc, "Связано с конкурентом", a.get("is_competitor_related"))
        _kv(doc, "Модель", a.get("model_used"))
        _kv(doc, "Дата анализа", a.get("analyzed_at"))
        if a.get("pain_points"):
            _heading(doc, "Боли аудитории", level + 2)
            for p in a["pain_points"]:
                _bullet(doc, p)
        if a.get("hooks"):
            _heading(doc, "Триггеры", level + 2)
            for h in a["hooks"]:
                _bullet(doc, h)
    else:
        doc.add_paragraph("AI-анализ: не выполнен.")

    doc.add_paragraph()


def save_report_docx(data: dict[str, Any], output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = build_report_docx(data)
    doc.save(output_path)
    return output_path
